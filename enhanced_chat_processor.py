import re
import io
import os
import hashlib
import datetime
from typing import List, Dict, Any, Optional
import time
import tempfile
import glob
import streamlit as st

# Try to import file handling libraries
try:
    import pandas as pd
except ImportError:
    pd = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx
except ImportError:
    docx = None

class EnhancedChatProcessor:
    """Enhanced processor for extracting and parsing chat transcripts from various file formats"""

    def extract_chats_from_file(self, uploaded_file):
        """
        Extract chats from uploaded file based on file type

        Args:
        uploaded_file: The uploaded file object

        Returns:
        List of chat dictionaries
        """
        try:
            file_extension = uploaded_file.name.split('.')[-1].lower()

            if file_extension == 'txt':
                return self._extract_from_txt(uploaded_file)
            elif file_extension == 'csv':
                if pd is None:
                    raise ImportError("pandas library is required for CSV processing. Install with: pip install pandas")
                return self._extract_from_csv(uploaded_file)
            elif file_extension == 'pdf':
                if PyPDF2 is None:
                    raise ImportError("PyPDF2 library is required for PDF processing. Install with: pip install PyPDF2")
                return self._extract_from_pdf(uploaded_file)
            elif file_extension == 'docx':
                if docx is None:
                    raise ImportError("python-docx library is required for DOCX processing. Install with: pip install python-docx")
                return self._extract_from_docx(uploaded_file)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
        except Exception as e:
            st.error(f"Error processing file: {str(e)}")
            return []

    def _extract_from_txt(self, uploaded_file):
        """Extract chats from a text file"""
        content = uploaded_file.getvalue().decode('utf-8')
        return self._split_text_into_chats(content)

    def _extract_from_csv(self, uploaded_file):
        """Extract chats from a CSV file"""
        # Read CSV into pandas DataFrame
        df = pd.read_csv(uploaded_file)

        # Check for common columns that might contain chat content
        content_columns = [col for col in df.columns if any(
            term in col.lower() for term in ['chat', 'message', 'content', 'text', 'transcript']
        )]

        if not content_columns:
            # If no obvious content column, use the first column containing text data
            for col in df.columns:
                if df[col].dtype == 'object':
                    content_columns = [col]
                    break

        # If still no content column found, use all columns
        if not content_columns:
            st.warning("Couldn't identify chat content columns. Using all columns.")
            content_columns = df.columns.tolist()

        # Join all content into a single string
        content = ""

        if len(content_columns) == 1:
            # Use entire CSV as one chat
            main_col = content_columns[0]
            content = "\n".join(df[main_col].astype(str).tolist())
        else:
            # Try to format as a dialogue
            for _, row in df.iterrows():
                for col in content_columns:
                    content += f"{col}: {row[col]}\n"
                content += "\n"

        return self._split_text_into_chats(content)

    def _extract_from_pdf(self, uploaded_file):
        """Extract chats from a PDF file"""
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.getvalue()))
        content = ""

        # Extract text from all pages
        for page_num in range(len(pdf_reader.pages)):
            content += pdf_reader.pages[page_num].extract_text() + "\n\n"

        return self._split_text_into_chats(content)

    def _extract_from_docx(self, uploaded_file):
        """Extract chats from a DOCX file"""
        doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
        
        # Process paragraphs with better structure preservation
        content_lines = []
        current_line = ""
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:  # Empty paragraph
                if current_line:
                    content_lines.append(current_line)
                    current_line = ""
                content_lines.append("")  # Preserve empty lines for structure
                continue
                
            # Check if this is a continuation of the previous line
            if text.startswith((':', '-', '...')) and current_line:
                current_line += text
            else:
                if current_line:
                    content_lines.append(current_line)
                current_line = text
        
        # Add the last line if any
        if current_line:
            content_lines.append(current_line)
            
        # Join lines with proper spacing    
        content = "\n".join(content_lines)
        
        # Add debug info
        #st.write("Extracted content structure:")
        #st.write("Number of paragraphs:", len(doc.paragraphs))
        #st.write("Number of content lines:", len(content_lines))
        #st.write("First few lines:", "\n".join(content_lines[:10]))
        
        # Process the content
        chats = self._split_text_into_chats(content)
        
        # Add debug info about extracted chats
        if not chats:
            st.warning("No chats were extracted. Content may not match expected format.")
            st.write("Content preview:")
            st.code(content[:500])  # Show first 500 chars
            
        return chats

    def _split_text_into_chats(self, text: str) -> List[Dict[str, Any]]:
        """
        Split a text containing multiple chat transcripts into individual chats.
        Enhanced to better handle various chat formats and language variations.
        """
        chats = []
        
        # Define minimum content requirements
        MIN_LINES = 3  # Minimum number of lines to consider as valid chat
        MIN_CHARS = 50  # Minimum number of characters to consider as valid chat
        
        # Focus on chat number patterns - match any number of digits
        chat_start_patterns = [
            r"(?:^|\s)Chat\s+(?:ID\s*:?\s*)?(\d+)(?!\d)",  # Match "Chat 01271153" (any number of digits)
            r"(?:^|\s)Chat\s*#\s*(\d+)(?!\d)",             # Match "Chat #01271153"
            r"^[A-Z]{2,}\s*[:-]?\s*Chat\s+(\d+)(?!\d)",    # Match "TH: Chat 01271153"
            r"^[A-Z]{2,}\s*[:-]?\s*Chat\s*#\s*(\d+)(?!\d)" # Match "TH: Chat #01271153"
        ]
        
        # Create combined pattern
        combined_pattern = '|'.join(f"(?:{pattern})" for pattern in chat_start_patterns)
        
        # Split by chat number patterns
        segments = re.split(f"({combined_pattern})", text, flags=re.IGNORECASE | re.MULTILINE)
        
        # Filter out empty segments
        segments = [segment for segment in segments if segment and not segment.isspace()]
        
        # Debug info for chat boundaries
        #st.write("Found chat boundaries:")
        #for segment in segments:
        #    if re.match(combined_pattern, segment, re.IGNORECASE):
        #        st.write(f"- {segment.strip()}")
        
        # Process segments  
        current_chat = ""
        chat_counter = 0
        is_header = False
        
        for i, segment in enumerate(segments):
            # Check if segment is a chat boundary
            if re.match(combined_pattern, segment, re.IGNORECASE):
                # Process previous chat if it exists and meets minimum requirements
                if current_chat.strip():
                    lines = [line for line in current_chat.split('\n') if line.strip()]
                    if len(lines) >= MIN_LINES and len(current_chat.strip()) >= MIN_CHARS:
                        chat_id = self._extract_chat_id(current_chat) or f"Chat_{chat_counter}"
                        timestamp = self._extract_timestamp(current_chat)
                        processed_content = self._clean_and_process_chat(current_chat)
                        
                        if processed_content:
                            chats.append({
                                'id': chat_id,
                                'content': current_chat,
                                'timestamp': timestamp or datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                                'processed_content': processed_content
                            })
                            chat_counter += 1
                
                current_chat = segment
                is_header = True
            else:
                # If this is content following a header, or we're still building the first chat
                if is_header or not current_chat:
                    current_chat += segment
                    is_header = False
                else:
                    # This is content not following a header, so it belongs to the current chat
                    current_chat += segment
        
        # Process final chat
        if current_chat.strip():
            lines = [line for line in current_chat.split('\n') if line.strip()]
            if len(lines) >= MIN_LINES and len(current_chat.strip()) >= MIN_CHARS:
                chat_id = self._extract_chat_id(current_chat) or f"Chat_{chat_counter}"
                timestamp = self._extract_timestamp(current_chat)
                processed_content = self._clean_and_process_chat(current_chat)
                
                if processed_content:
                    chats.append({
                        'id': chat_id,
                        'content': current_chat,
                        'timestamp': timestamp or datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        'processed_content': processed_content
                    })
        
        return chats

    def _extract_chat_id(self, chat_text: str) -> Optional[str]:
        """Extract chat ID from chat text."""
        patterns = [
            r"(?:^|\s)Chat\s+(?:ID\s*:?\s*)?(\d+)(?!\d)",  # Match "Chat 01271153"
            r"(?:^|\s)Chat\s*#\s*(\d+)(?!\d)",             # Match "Chat #01271153"
            r"^[A-Z]{2,}\s*[:-]?\s*Chat\s+(\d+)(?!\d)",    # Match "TH: Chat 01271153"
            r"^[A-Z]{2,}\s*[:-]?\s*Chat\s*#\s*(\d+)(?!\d)" # Match "TH: Chat #01271153"
        ]
    
        for pattern in patterns:
            match = re.search(pattern, chat_text, re.IGNORECASE)
            if match:
                return f"Chat_{match.group(1)}"
        
        # If no valid chat ID found, generate a hash from the content
        content_sample = chat_text[:100].strip()
        if content_sample:
            hash_obj = hashlib.md5(content_sample.encode())
            return f"Chat_{hash_obj.hexdigest()[:8]}"
        
        return None

    def _extract_timestamp(self, chat_text: str) -> Optional[str]:
        """Extract timestamp from chat text."""
        timestamp_patterns = [
            r"Chat Started:\s*([^,\n]+(?:\s+\d{4})?)",
            r"Session Started:\s*([^,\n]+(?:\s+\d{4})?)",
            r"Conversation Started:\s*([^,\n]+(?:\s+\d{4})?)"
        ]

        for pattern in timestamp_patterns:
            match = re.search(pattern, chat_text, re.IGNORECASE)
            if match:
                try:
                    timestamp_str = match.group(1).strip()
                    # Try to parse the timestamp - handle various formats
                    for fmt in [
                        "%A, %B %d, %Y, %H:%M:%S",
                        "%A, %B %d, %Y %H:%M:%S",
                        "%Y-%m-%d %H:%M:%S"
                    ]:
                        try:
                            parsed_time = datetime.datetime.strptime(timestamp_str, fmt)
                            return parsed_time.strftime("%Y-%m-%d %H:%M:%S")
                        except ValueError:
                            continue
                except Exception:
                    pass
        return None

    def _clean_and_process_chat(self, chat_text: str) -> str:
        """Clean and format chat content for analysis."""
        lines = chat_text.split('\n')
        processed_lines = []
        has_valid_content = False
        speaker_patterns = {
            'customer': [
                r'^(?:Customer|Client|Visitor|User|Guest)',
                r'^\(\d+[ms]\)\s*(?:Customer|Client|Visitor|User|Guest)',
                r'^.*?(?:Customer|Client|Visitor|User|Guest)\s*:'
            ],
            'agent': [
                r'^(?:Agent|Support|Assistant|Representative|Bot|Chatbot|Pepper)',
                r'^\(\d+[ms]\)\s*(?:Agent|Support|Assistant|Representative|Bot|Chatbot|Pepper)',
                r'^.*?(?:Agent|Support|Assistant|Representative|Bot|Chatbot|Pepper)\s*:'
            ]
        }

        # Simplified skip patterns - only skip non-conversation content
        skip_patterns = [
            r'^\s*\{ChatWindowButton:',
            r'^\s*Page \d+ of \d+',
            r'^\s*Report generated',
            r'^\s*[-=*_]{3,}\s*$'  # Separator lines
        ]

        skip_pattern = '|'.join(skip_patterns)
        has_conversation = False

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Skip non-conversation content
            if re.match(skip_pattern, line, re.IGNORECASE):
                continue

            # Check for any type of message
            is_message = False
            # Check for customer messages
            for pattern in speaker_patterns['customer']:
                if re.match(pattern, line, re.IGNORECASE):
                    is_message = True
                    has_valid_content = True
                    has_conversation = True
                    # Clean up the line
                    line = re.sub(r'^\(\d+[ms]\)\s*', '', line)  # Remove timestamp
                    line = re.sub(r'^.*?(?:Customer|Client|Visitor|User|Guest)\s*:', 'Customer:', line, flags=re.IGNORECASE)
                    break

            # Check for agent/bot messages
            if not is_message:
                for pattern in speaker_patterns['agent']:
                    if re.match(pattern, line, re.IGNORECASE):
                        is_message = True
                        has_valid_content = True
                        has_conversation = True
                        # Clean up the line
                        line = re.sub(r'^\(\d+[ms]\)\s*', '', line)  # Remove timestamp
                        # Preserve the original speaker type (Bot/Agent)
                        if re.match(r'^.*?(Bot|Chatbot|Pepper)\s*:', line, re.IGNORECASE):
                            line = re.sub(r'^.*?(Bot|Chatbot|Pepper)\s*:', r'\1:', line, flags=re.IGNORECASE)
                        else:
                            line = re.sub(r'^.*?(?:Agent|Support|Assistant|Representative)\s*:', 'Agent:', line, flags=re.IGNORECASE)
                        break

            if is_message:
                processed_lines.append(line)
            elif line.strip():  # Add non-empty lines that might be continuation of messages
                processed_lines.append(line)

        # Return content if it has any conversation
        if has_valid_content and has_conversation:
            return '\n'.join(processed_lines)
        return ''

def cleanup_session_state():
    """Clean up old session data"""
    max_age = 3600  # 1 hour
    now = time.time()
    for key in list(st.session_state.keys()):
        if key.startswith('temp_') and now - st.session_state[f"{key}_timestamp"] > max_age:
            del st.session_state[key]

class RateLimiter:
    def __init__(self, calls_per_minute=60):
        self.calls_per_minute = calls_per_minute
        self.calls = []
    
    def wait_if_needed(self):
        now = time.time()
        self.calls = [t for t in self.calls if now - t < 60]
        if len(self.calls) >= self.calls_per_minute:
            time.sleep(60 - (now - self.calls[0]))
        self.calls.append(now)

def cleanup_temp_files():
    """Clean up temporary audio files"""
    temp_dir = tempfile.gettempdir()
    pattern = "qa_engine_*.wav"
    for f in glob.glob(os.path.join(temp_dir, pattern)):
        try:
            os.remove(f)
        except OSError:
            pass

@st.cache_data(ttl=3600)
def process_file_cached(file_content, file_type):
    """Cache file processing results"""
    processor = EnhancedChatProcessor()
    return processor.extract_chats_from_file(file_content)